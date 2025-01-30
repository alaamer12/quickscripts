from docx import Document
from docx.shared import Pt

# Define the content for each component
components = {
    "Resources Document": {
        "Purpose": "The Resources Document outlines the materials and assets required for a project or task.",
        "Contribution": "Lists essential resources to ensure availability and effective project execution.",
        "Overall Success": "Supports efficient project management and resource allocation."
    },
    "Action Plan": {
        "Purpose": "The Action Plan details the steps and strategies for achieving specific goals or objectives.",
        "Contribution": "Provides a clear roadmap for task execution and goal attainment.",
        "Overall Success": "Ensures organized and strategic approach towards achieving goals."
    },
    "Annual Plan": {
        "Purpose": "The Annual Plan outlines the organization’s objectives and strategies for the year.",
        "Contribution": "Guides long-term planning and sets the direction for the organization’s annual activities.",
        "Overall Success": "Supports strategic alignment and goal setting for the year."
    },
    "Business Plan": {
        "Purpose": "The Business Plan outlines the organization’s objectives, strategies, and financial projections.",
        "Contribution": "Provides a comprehensive overview of the business strategy and financial outlook.",
        "Overall Success": "Supports strategic planning and attracts potential investors or stakeholders."
    },
    "Decision Matrix": {
        "Purpose": "The Decision Matrix evaluates and prioritizes different options based on predefined criteria.",
        "Contribution": "Assists in making informed decisions by comparing and scoring options.",
        "Overall Success": "Facilitates objective decision-making and ensures alignment with goals."
    },
    "Decision Tree [Strategy]": {
        "Purpose": "The Decision Tree [Strategy] visually represents different decision paths and their possible outcomes.",
        "Contribution": "Helps in understanding the potential impacts of different decisions and strategies.",
        "Overall Success": "Supports strategic planning by illustrating decision-making processes and consequences."
    },
    "Promotion Event Document": {
        "Purpose": "The Promotion Event Document outlines the details and plan for a promotional event.",
        "Contribution": "Provides a structured approach to organizing and executing promotional activities.",
        "Overall Success": "Enhances the effectiveness of promotional events and improves engagement."
    },
    "User Guide": {
        "Purpose": "The User Guide provides instructions and information on how to use a product or service.",
        "Contribution": "Helps users understand and effectively utilize the product or service.",
        "Overall Success": "Improves user experience and reduces support requests by offering clear guidance."
    },
    "Business Contingency Plan": {
        "Purpose": "The Business Contingency Plan outlines procedures for handling unexpected disruptions or emergencies.",
        "Contribution": "Ensures that critical business functions can continue during and after a crisis.",
        "Overall Success": "Supports organizational resilience and minimizes the impact of disruptions."
    },
    "Version Plan": {
        "Purpose": "The Version Plan details the management and release schedule for different software versions.",
        "Contribution": "Organizes versioning and updates to ensure smooth software evolution.",
        "Overall Success": "Facilitates effective version control and release management."
    },
    "Release Plan": {
        "Purpose": "The Release Plan outlines the strategy and schedule for deploying software or product releases.",
        "Contribution": "Coordinates release activities to ensure successful and timely launches.",
        "Overall Success": "Supports organized product releases and minimizes deployment issues."
    },
    "Marketing Plan": {
        "Purpose": "The Marketing Plan defines strategies and tactics for promoting products or services.",
        "Contribution": "Guides marketing efforts to attract and retain customers.",
        "Overall Success": "Enhances brand visibility and drives customer engagement through targeted marketing."
    },
    "Financial Plan": {
        "Purpose": "The Financial Plan outlines the organization’s financial goals, projections, and strategies.",
        "Contribution": "Provides a roadmap for managing financial resources and achieving fiscal objectives.",
        "Overall Success": "Supports sound financial management and aligns with organizational goals."
    },
    "Dashboard Analysis Documents": {
        "Purpose": "The Dashboard Analysis Documents provide insights and visualizations of key performance metrics.",
        "Contribution": "Offers a comprehensive view of performance and trends for informed decision-making.",
        "Overall Success": "Enhances data-driven decision-making and performance tracking."
    },
    "Simple Job Description": {
        "Purpose": "The Simple Job Description outlines the primary responsibilities and requirements for a role.",
        "Contribution": "Provides a clear overview of job expectations and qualifications.",
        "Overall Success": "Supports effective recruitment and role clarity."
    },
    "Future Plan": {
        "Purpose": "The Future Plan details the organization’s long-term goals and strategies for growth.",
        "Contribution": "Guides future development and strategic direction.",
        "Overall Success": "Supports long-term planning and aligns with organizational vision."
    },
    "Team Work Plan": {
        "Purpose": "The Team Work Plan outlines tasks, responsibilities, and timelines for team members.",
        "Contribution": "Facilitates team coordination and task management.",
        "Overall Success": "Enhances team productivity and ensures alignment with project goals."
    },
    "Legal Plan": {
        "Purpose": "The Legal Plan outlines legal requirements and strategies for managing legal risks.",
        "Contribution": "Ensures compliance with regulations and protects the organization’s interests.",
        "Overall Success": "Supports legal and regulatory adherence and risk management."
    },
    "Deployment Plan": {
        "Purpose": "The Deployment Plan details the steps and procedures for releasing software or products.",
        "Contribution": "Coordinates deployment activities to ensure smooth releases.",
        "Overall Success": "Facilitates effective product launches and minimizes deployment issues."
    },
    "Regular Exporting Plan": {
        "Purpose": "The Regular Exporting Plan outlines procedures for exporting data or products on a routine basis.",
        "Contribution": "Ensures consistent and reliable exporting processes.",
        "Overall Success": "Supports regular data management and export operations."
    },
    "Development Plan": {
        "Purpose": "The Development Plan outlines the approach and phases for developing software or products.",
        "Contribution": "Provides a structured framework for development activities.",
        "Overall Success": "Supports organized and efficient development processes."
    },
    "Staging Plan": {
        "Purpose": "The Staging Plan details the procedures for preparing a product or software for final release.",
        "Contribution": "Ensures that the product is thoroughly tested and validated before deployment.",
        "Overall Success": "Facilitates a smooth transition from development to production."
    },
    "Production Plan": {
        "Purpose": "The Production Plan outlines the processes and resources required for manufacturing or producing a product.",
        "Contribution": "Guides production activities and resource allocation.",
        "Overall Success": "Ensures efficient and effective production operations."
    },
    "Canary Plan": {
        "Purpose": "The Canary Plan involves releasing new features or updates to a small subset of users before full deployment.",
        "Contribution": "Allows for testing in a live environment with minimal risk.",
        "Overall Success": "Identifies issues early and ensures a smoother full release."
    },
    "Targeting Plan": {
        "Purpose": "The Targeting Plan outlines strategies for reaching specific audiences or markets.",
        "Contribution": "Guides marketing and outreach efforts to focus on key target groups.",
        "Overall Success": "Enhances marketing effectiveness and audience engagement."
    },
    "Event Plan": {
        "Purpose": "The Event Plan details the organization and execution of events.",
        "Contribution": "Provides a structured approach to event planning and management.",
        "Overall Success": "Ensures successful event execution and engagement."
    },
    "Pre-awareness Plan": {
        "Purpose": "The Pre-awareness Plan outlines strategies for generating interest and anticipation before a product or event launch.",
        "Contribution": "Builds awareness and excitement prior to the official release.",
        "Overall Success": "Enhances pre-launch visibility and engagement."
    },
    "Awareness Plan": {
        "Purpose": "The Awareness Plan details strategies for increasing visibility and awareness of a product or service.",
        "Contribution": "Guides efforts to promote and publicize the product or service.",
        "Overall Success": "Boosts product visibility and attracts potential customers."
    },
    "Social Pages Content Plan": {
        "Purpose": "The Social Pages Content Plan outlines the content strategy for social media channels.",
        "Contribution": "Provides a framework for creating and scheduling social media content.",
        "Overall Success": "Enhances social media engagement and brand presence."
    },
    "Contact List": {
        "Purpose": "The Contact List includes key contact information for individuals or organizations.",
        "Contribution": "Facilitates communication and networking.",
        "Overall Success": "Supports effective contact management and outreach."
    },
    "Stakeholders Contact List": {
        "Purpose": "The Stakeholders Contact List contains contact details for project stakeholders.",
        "Contribution": "Provides essential information for stakeholder communication and engagement.",
        "Overall Success": "Enhances stakeholder management and collaboration."
    },
    "Shareholders List": {
        "Purpose": "The Shareholders List includes information on individuals or entities owning shares in the organization.",
        "Contribution": "Facilitates communication and reporting to shareholders.",
        "Overall Success": "Supports shareholder relations and transparency."
    },
    "COCOMO": {
        "Purpose": "COCOMO (Constructive Cost Model) is a model used for estimating software development costs and project timelines.",
        "Contribution": "Provides cost and schedule estimates based on project size and complexity.",
        "Overall Success": "Helps in budgeting and planning by offering a data-driven approach to project management."
    },
    "Economic Feasibility Study": {
        "Purpose": "The Economic Feasibility Study assesses the financial viability of a project or investment.",
        "Contribution": "Evaluates the cost-benefit ratio and financial impact.",
        "Overall Success": "Supports decision-making by providing insights into the financial implications of a project."
    },
    "System Analysis": {
        "Purpose": "System Analysis involves studying and understanding the components and interactions within a system.",
        "Contribution": "Identifies system requirements and issues for improvement.",
        "Overall Success": "Supports effective system design and optimization."
    },
    "Backlog": {
        "Purpose": "The Backlog is a prioritized list of tasks or features to be addressed in a project.",
        "Contribution": "Organizes work items for future development or resolution.",
        "Overall Success": "Ensures that important tasks are addressed in a structured manner."
    },
    "Story Map": {
        "Purpose": "The Story Map visualizes user stories and their relationships to provide a comprehensive view of the product backlog.",
        "Contribution": "Helps in understanding user needs and prioritizing features.",
        "Overall Success": "Supports effective backlog management and product planning."
    },
    "Deployment Metadata": {
        "Purpose": "Deployment Metadata includes details about the deployment process and configurations.",
        "Contribution": "Provides information for managing and tracking deployment activities.",
        "Overall Success": "Supports effective deployment and configuration management."
    },
    "Patch Release Plan": {
        "Purpose": "The Patch Release Plan outlines the process for releasing minor updates or fixes to a product.",
        "Contribution": "Ensures timely and effective deployment of patches.",
        "Overall Success": "Facilitates ongoing product maintenance and improvement."
    },
    "Minor Release Plan": {
        "Purpose": "The Minor Release Plan details the approach for releasing incremental updates to a product.",
        "Contribution": "Coordinates minor updates and enhancements.",
        "Overall Success": "Supports product evolution and feature improvements."
    },
    "Major Release Plan": {
        "Purpose": "The Major Release Plan outlines the strategy and schedule for significant product releases.",
        "Contribution": "Manages large-scale updates and new feature deployments.",
        "Overall Success": "Ensures successful implementation of major updates and enhancements."
    },
    "Alpha Release Plan": {
        "Purpose": "The Alpha Release Plan details the process for releasing early versions of a product to a select group of users.",
        "Contribution": "Facilitates initial testing and feedback collection.",
        "Overall Success": "Supports early-stage product validation and refinement."
    },
    "Beta Release Plan": {
        "Purpose": "The Beta Release Plan outlines the strategy for releasing a product to a broader audience for testing.",
        "Contribution": "Gathers feedback and identifies issues before the final release.",
        "Overall Success": "Enhances product quality through user testing and feedback."
    },
    "Canary Release Plan": {
        "Purpose": "The Canary Release Plan involves gradually rolling out new features to a small subset of users.",
        "Contribution": "Allows for testing in a live environment with reduced risk.",
        "Overall Success": "Supports gradual feature deployment and issue detection."
    },
    "Bug Report": {
        "Purpose": "The Bug Report documents issues or defects found in software or products.",
        "Contribution": "Provides details for developers to address and fix issues.",
        "Overall Success": "Supports product quality and user satisfaction by facilitating timely bug resolution."
    },
    "Release Map": {
        "Purpose": "The Release Map visualizes the schedule and process for upcoming product releases.",
        "Contribution": "Provides a clear overview of release timelines and milestones.",
        "Overall Success": "Supports effective release planning and tracking."
    },
    "Branding Guidelines": {
        "Purpose": "The Branding Guidelines document outlines the standards for brand presentation and identity.",
        "Contribution": "Ensures consistent and professional brand representation.",
        "Overall Success": "Supports brand integrity and recognition."
    },
    "Project Charter": {
        "Purpose": "The Project Charter defines the project’s objectives, scope, and stakeholders.",
        "Contribution": "Provides a foundational document for project initiation and planning.",
        "Overall Success": "Supports project alignment and stakeholder understanding."
    },
    "SRS": {
        "Purpose": "The Software Requirements Specification (SRS) document details the functional and non-functional requirements for a software system.",
        "Contribution": "Provides a comprehensive description of system requirements for development.",
        "Overall Success": "Ensures clear and accurate requirements for successful software development."
    },
    "PERT": {
        "Purpose": "The Program Evaluation and Review Technique (PERT) is used for project scheduling and management.",
        "Contribution": "Helps in estimating project duration and managing tasks.",
        "Overall Success": "Supports effective project planning and control."
    },
    "Sprints": {
        "Purpose": "Sprints are time-boxed periods in which specific work is completed and reviewed.",
        "Contribution": "Organizes work into manageable intervals for iterative development.",
        "Overall Success": "Enhances project agility and iterative progress."
    },
    "LICENSE": {
        "Purpose": "The LICENSE document specifies the terms under which software or content can be used and distributed.",
        "Contribution": "Provides legal guidelines and permissions for use.",
        "Overall Success": "Ensures compliance with licensing terms and protects intellectual property."
    },
    "Changelog": {
        "Purpose": "The Changelog records changes, updates, and fixes made to a product or software.",
        "Contribution": "Provides a historical record of changes for users and developers.",
        "Overall Success": "Supports transparency and helps in tracking product evolution."
    },
    "Terms of Service [ToS]": {
        "Purpose": "The Terms of Service document outlines the rules and conditions for using a service or product.",
        "Contribution": "Sets expectations and legal terms for users.",
        "Overall Success": "Supports legal compliance and user understanding."
    },
    "Disclaimer": {
        "Purpose": "The Disclaimer provides information on limitations and responsibilities related to the use of a service or product.",
        "Contribution": "Clarifies liability and limitations of use.",
        "Overall Success": "Helps in managing legal risks and user expectations."
    },
    "Policies": {
        "Purpose": "The Policies document outlines organizational rules and procedures.",
        "Contribution": "Provides guidance and standards for organizational operations.",
        "Overall Success": "Ensures consistency and compliance with organizational standards."
    },
    "API Documentation": {
        "Purpose": "API Documentation provides information on how to use and integrate with an API.",
        "Contribution": "Offers detailed instructions and examples for developers.",
        "Overall Success": "Enhances developer experience and supports effective API usage."
    },
    "BCP": {
        "Purpose": "The Business Continuity Plan (BCP) outlines procedures for maintaining business operations during disruptions.",
        "Contribution": "Ensures continuity of critical business functions.",
        "Overall Success": "Supports organizational resilience and crisis management."
    },
    "Asset Tracking": {
        "Purpose": "The Asset Tracking document details procedures for managing and monitoring organizational assets.",
        "Contribution": "Provides a systematic approach to asset management.",
        "Overall Success": "Enhances asset visibility and control."
    },
    "Compliance": {
        "Purpose": "The Compliance document outlines adherence to regulatory and legal requirements.",
        "Contribution": "Ensures that organizational practices meet legal standards.",
        "Overall Success": "Supports regulatory adherence and risk management."
    },
    "Information Security Procedures": {
        "Purpose": "The Information Security Procedures document details practices for safeguarding organizational data.",
        "Contribution": "Provides guidelines for protecting information and preventing breaches.",
        "Overall Success": "Enhances data security and compliance with security policies."
    },
    "Centers Contact List": {
        "Purpose": "The Centers Contact List includes contact information for various organizational centers.",
        "Contribution": "Facilitates communication and coordination with centers.",
        "Overall Success": "Supports effective outreach and interaction with organizational centers."
    },
    "Future Features": {
        "Purpose": "The Future Features document outlines planned enhancements and features for a product.",
        "Contribution": "Provides a roadmap for future development and improvements.",
        "Overall Success": "Guides product evolution and aligns with user needs."
    },
    "Incident Response Plan [IRP]": {
        "Purpose": "The Incident Response Plan details procedures for responding to and managing incidents.",
        "Contribution": "Ensures a structured approach to incident management and resolution.",
        "Overall Success": "Supports effective incident handling and minimizes impact."
    },
    "Information Security Policy": {
        "Purpose": "The Information Security Policy defines the organization’s approach to information security.",
        "Contribution": "Establishes standards and practices for protecting organizational data.",
        "Overall Success": "Ensures data protection and compliance with security requirements."
    },
    "Risk Assessment Report": {
        "Purpose": "The Risk Assessment Report evaluates potential risks and their impact on the organization.",
        "Contribution": "Identifies vulnerabilities and informs risk management strategies.",
        "Overall Success": "Supports proactive risk management and minimizes potential negative impacts."
    },
    "Code of Conduct": {
        "Purpose": "The Code of Conduct sets forth the ethical and behavioral standards expected of employees and stakeholders.",
        "Contribution": "Promotes a positive and professional work environment.",
        "Overall Success": "Ensures consistent ethical behavior and supports organizational values."
    },
    "Contribution": {
        "Purpose": "The Contribution document acknowledges and details the contributions of individuals or teams to a project or organization.",
        "Contribution": "Recognizes efforts and provides credit for achievements.",
        "Overall Success": "Enhances motivation and fosters a culture of appreciation and recognition."
    },
    "Contributors": {
        "Purpose": "The Contributors document lists individuals or teams who have contributed to a project or initiative.",
        "Contribution": "Provides recognition and thanks for contributions.",
        "Overall Success": "Supports collaboration and acknowledges valuable input."
    },
    "SQL File": {
        "Purpose": "The SQL File contains database queries and scripts for managing and manipulating data.",
        "Contribution": "Provides the means for database operations and data management.",
        "Overall Success": "Supports database functionality and data processing."
    },
    "Testing Strategy": {
        "Purpose": "The Testing Strategy outlines the approach and methods for testing a product or system.",
        "Contribution": "Provides a framework for ensuring product quality and functionality.",
        "Overall Success": "Enhances testing effectiveness and identifies issues before release."
    },
    "User Guides": {
        "Purpose": "User Guides provide comprehensive instructions and information for using a product or service.",
        "Contribution": "Helps users understand and effectively utilize the product or service.",
        "Overall Success": "Improves user experience and reduces support needs."
    },
    "Architecture Decision Records [ADR]": {
        "Purpose": "Architecture Decision Records document key decisions made during the software design process.",
        "Contribution": "Provides a record of architectural choices and rationale.",
        "Overall Success": "Supports transparency and understanding of architectural decisions."
    },
    "Project [Product] Roadmap": {
        "Purpose": "The Project [Product] Roadmap outlines the planned development and release schedule for a project or product.",
        "Contribution": "Provides a visual representation of project milestones and goals.",
        "Overall Success": "Supports strategic planning and project management."
    },
    "Post-Mortem Report": {
        "Purpose": "The Post-Mortem Report analyzes and reviews a completed project to assess successes and failures.",
        "Contribution": "Provides insights and lessons learned for future projects.",
        "Overall Success": "Supports continuous improvement and project learning."
    },
    "README": {
        "Purpose": "The README document provides an overview and instructions for using a project or repository.",
        "Contribution": "Offers essential information and guidance for users and developers.",
        "Overall Success": "Enhances understanding and usability of the project."
    },
    "AUTHORS": {
        "Purpose": "The AUTHORS document lists individuals who have contributed to the creation or development of a project.",
        "Contribution": "Provides recognition and credit to contributors.",
        "Overall Success": "Supports acknowledgment and appreciation of contributions."
    },
    "ERD Diagram": {
        "Purpose": "The Entity-Relationship Diagram (ERD) visualizes the relationships between entities in a database.",
        "Contribution": "Provides a graphical representation of database structure.",
        "Overall Success": "Supports database design and understanding of data relationships."
    },
    "Use-case Diagram": {
        "Purpose": "The Use-case Diagram illustrates the interactions between users and a system.",
        "Contribution": "Helps in understanding user requirements and system functionality.",
        "Overall Success": "Supports requirement analysis and system design."
    },
    "High Level Application Diagram": {
        "Purpose": "The High Level Application Diagram provides an overview of the system architecture and components.",
        "Contribution": "Offers a broad view of system design and interactions.",
        "Overall Success": "Supports architectural planning and communication."
    },
    "CPM Diagram": {
        "Purpose": "The Critical Path Method (CPM) Diagram visualizes project tasks and their dependencies.",
        "Contribution": "Helps in identifying the critical path and managing project timelines.",
        "Overall Success": "Supports project scheduling and management."
    },
    "Gantt Diagram": {
        "Purpose": "The Gantt Diagram provides a timeline view of project tasks and milestones.",
        "Contribution": "Offers a visual representation of project schedules and progress.",
        "Overall Success": "Supports project planning and tracking."
    },
    "API Hierarchy Diagram": {
        "Purpose": "The API Hierarchy Diagram visualizes the structure and relationships of API components.",
        "Contribution": "Provides an overview of API design and organization.",
        "Overall Success": "Supports API development and understanding."
    },
    "API State Diagram": {
        "Purpose": "The API State Diagram illustrates the different states and transitions of an API.",
        "Contribution": "Helps in understanding API behavior and state management.",
        "Overall Success": "Supports API design and troubleshooting."
    },
    "Workflow Diagram": {
        "Purpose": "The Workflow Diagram visualizes the flow of tasks and processes within a system.",
        "Contribution": "Provides a graphical representation of process sequences and interactions.",
        "Overall Success": "Supports process design and optimization."
    },
    "CSD Diagram": {
        "Purpose": "The Component-Subsystem Diagram (CSD) visualizes the components and subsystems of a system.",
        "Contribution": "Offers an overview of system architecture and component interactions.",
        "Overall Success": "Supports system design and analysis."
    },
    "DFD Level 0 Diagram": {
        "Purpose": "The Data Flow Diagram (DFD) Level 0 provides a high-level view of system processes and data flows.",
        "Contribution": "Offers a broad understanding of system interactions and data movement.",
        "Overall Success": "Supports system analysis and design."
    },
    "DFD Level 1 Diagram": {
        "Purpose": "The Data Flow Diagram (DFD) Level 1 provides a more detailed view of system processes and data flows.",
        "Contribution": "Offers an in-depth understanding of system processes and data interactions.",
        "Overall Success": "Supports detailed system design and analysis."
    },
    "WBS Diagram": {
        "Purpose": "The Work Breakdown Structure (WBS) Diagram visualizes the breakdown of project tasks and deliverables.",
        "Contribution": "Provides a hierarchical view of project components and tasks.",
        "Overall Success": "Supports project planning and management."
    },
    "Feature-based Roadmap Diagram": {
        "Purpose": "The Feature-based Roadmap Diagram outlines the development and release schedule of product features.",
        "Contribution": "Provides a visual representation of feature development and timelines.",
        "Overall Success": "Supports feature planning and prioritization."
    },
    "Goal-based Roadmap Diagram": {
        "Purpose": "The Goal-based Roadmap Diagram visualizes the goals and milestones for a project or product.",
        "Contribution": "Provides a structured view of project objectives and progress.",
        "Overall Success": "Supports goal alignment and tracking."
    },
    "Timeline-based Roadmap Diagram": {
        "Purpose": "The Timeline-based Roadmap Diagram provides a chronological view of project milestones and activities.",
        "Contribution": "Offers a visual representation of project timelines and deadlines.",
        "Overall Success": "Supports project scheduling and tracking."
    },
    "BPMN Diagram": {
        "Purpose": "The Business Process Model and Notation (BPMN) Diagram visualizes business processes and workflows.",
        "Contribution": "Provides a standard notation for documenting business processes.",
        "Overall Success": "Supports process analysis and improvement."
    },
    "ERD Mapping Diagram": {
        "Purpose": "The ERD Mapping Diagram visualizes the mapping between different data entities and relationships.",
        "Contribution": "Offers a detailed view of data relationships and structure.",
        "Overall Success": "Supports database design and data management."
    }
}


def add_bullet_paragraph(doc, text, level=0):
    """Add a bullet point paragraph with indentation to the document."""
    para = doc.add_paragraph(text, style='List Bullet')
    # Adjust indentation
    para.paragraph_format.left_indent = Pt(12 * level)
    para.paragraph_format.space_after = Pt(6)  # Add some space after each bullet
    return para

def create_document(components, file_name="Components_Documentation.docx"):
    doc = Document()

    # Set document title
    title = doc.add_heading('Component Documentation', level=1)
    run = title.runs[0]
    run.font.size = Pt(24)

    for component, details in components.items():
        # Add component title
        comp_heading = doc.add_heading(component, level=2)
        comp_run = comp_heading.runs[0]
        comp_run.font.size = Pt(18)

        # Add purpose
        doc.add_heading('Purpose', level=3).runs[0].font.size = Pt(16)
        purpose_text = details['Purpose'].split('\n')
        for line in purpose_text:
            add_bullet_paragraph(doc, line, level=1)

        # Add contribution
        doc.add_heading('Contribution', level=3).runs[0].font.size = Pt(16)
        contribution_text = details['Contribution'].split('\n')
        for line in contribution_text:
            add_bullet_paragraph(doc, line, level=1)

        # Add overall success
        doc.add_heading('Overall Success', level=3).runs[0].font.size = Pt(16)
        success_text = details['Overall Success'].split('\n')
        for line in success_text:
            add_bullet_paragraph(doc, line, level=1)

        # Add space between components
        doc.add_paragraph("\n")

    # Save the document
    doc.save(file_name)
    print(f"Document '{file_name}' has been created.")
# Create the document
create_document(components)
