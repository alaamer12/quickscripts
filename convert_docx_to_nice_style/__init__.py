from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from typing import Optional, Tuple, List
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table
from rich import print as rprint
from rich.traceback import install

# Install rich traceback handler
install(show_locals=True)

# Initialize Rich console
console = Console()

# Modern color palette
COLORS = {
    'primary': (41, 128, 185),  # Blue
    'secondary': (52, 73, 94),  # Dark Gray
    'accent1': (46, 204, 113),  # Green
    'accent2': (155, 89, 182),  # Purple
    'text': (44, 62, 80),  # Dark Gray
    'light_bg': (236, 240, 241),  # Light Gray
}


def validate_file_path(file_path: str) -> Tuple[bool, str]:
    """Validate if file path exists and is a .docx file"""
    if not file_path:
        return False, "File path cannot be empty"
    if not file_path.endswith('.docx'):
        return False, "File must be a .docx document"
    if not os.path.exists(file_path):
        return False, f"File not found: {file_path}"
    return True, ""


def create_paragraph_style(styles, name: str, font_size: int, font_color: Tuple[int, int, int],
                           bold: bool = False, italic: bool = False) -> Optional[object]:
    """Helper function to create paragraph styles with error handling"""
    try:
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]

        font = style.font
        font.name = 'Calibri'
        font.size = Pt(max(8, min(font_size, 72)))  # Limit font size between 8 and 72
        font.color.rgb = RGBColor(*[max(0, min(x, 255)) for x in font_color])  # Ensure RGB values are valid
        font.bold = bold
        font.italic = italic
        return style
    except Exception as e:
        console.print(f"[red]Error creating style '{name}':[/red] {str(e)}")
        return None


def create_table_style(table) -> None:
    """Apply modern styling to a table with error handling"""
    try:
        # Set table properties
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Define cell shading XML
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ECEFF1"/>')

        # Apply styles to each cell
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                try:
                    # Header row
                    if row_idx == 0:
                        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>'))
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold = True
                    # Alternating rows
                    elif row_idx % 2 == 1:
                        cell._tc.get_or_add_tcPr().append(shading_elm)

                    # Vertical alignment
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Add padding
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(6)
                        paragraph.paragraph_format.space_after = Pt(6)
                except Exception as cell_error:
                    console.print(f"[red]Warning: Error styling cell in row {row_idx}:[/red] {str(cell_error)}")
                    continue
    except Exception as e:
        console.print(f"[red]Warning: Error applying table style:[/red] {str(e)}")


def process_paragraph_text(text: str) -> Tuple[bool, str]:
    """Clean and validate paragraph text"""
    if not text:
        return False, ""

    # Remove excessive whitespace
    text = ' '.join(text.split())

    # Skip unnecessary content
    skip_starts = ('chatgpt', 'you', 'certainly', 'but i think',
                   'i have always', 'here\'s', "here's")
    if text.lower().startswith(skip_starts):
        return False, ""

    return True, text


def is_role_title(text: str) -> bool:
    """Check if text represents a role title"""
    role_keywords = {'Officer', 'President', 'CEO', 'CTO', 'CIO', 'CSO',
                     'CFO', 'COO', 'CMO', 'CHRO', 'Director', 'Head',
                     'Manager', 'Lead'}

    return any(keyword in text for keyword in role_keywords)


def clean_and_style_document(input_path: str, output_path: str) -> bool:
    """Process a Word document and create a styled version"""
    # Validate input file
    is_valid, error_msg = validate_file_path(input_path)
    if not is_valid:
        console.print(f"[red]Error:[/red] {error_msg}")
        return False

    # Validate output path
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception as e:
            console.print(f"[red]Error creating output directory:[/red] {str(e)}")
            return False

    with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
    ) as progress:
        try:
            task_read = progress.add_task("[cyan]Reading document...", total=None)
            original_doc = Document(input_path)
            progress.update(task_read, completed=True)
        except Exception as e:
            console.print(f"[red]Error reading document:[/red] {str(e)}")
            return False

        # Create new document
        progress.add_task("[cyan]Creating new document...", total=None)
        new_doc = Document()

        # Process document with progress tracking
        task_process = progress.add_task("[cyan]Processing document...", total=100)

        # Set up page margins
        for section in new_doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
        progress.update(task_process, advance=10)

        # Create styles
        styles = new_doc.styles
        progress.update(task_process, advance=10)

        # Define styles with modern color scheme
        style_definitions = [
            ('Title Style', 28, COLORS['primary'], True, False),
            ('Subtitle Style', 16, COLORS['secondary'], False, True),
            ('Heading Style', 20, COLORS['primary'], True, False),
            ('Role Style', 16, COLORS['accent1'], True, False),
            ('Purpose Style', 11, COLORS['text'], False, False)
        ]

        created_styles = {}
        for style_name, size, color, bold, italic in style_definitions:
            style = create_paragraph_style(styles, style_name, size, color, bold, italic)
            if style:
                created_styles[style_name] = style

                # Apply specific formatting
                if style_name == 'Title Style':
                    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    style.paragraph_format.space_after = Pt(24)
                elif style_name == 'Subtitle Style':
                    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    style.paragraph_format.space_after = Pt(24)
                elif style_name == 'Heading Style':
                    style.paragraph_format.space_before = Pt(18)
                    style.paragraph_format.space_after = Pt(12)
                    style.paragraph_format.keep_with_next = True
                elif style_name == 'Role Style':
                    style.paragraph_format.left_indent = Inches(0.25)
                    style.paragraph_format.space_before = Pt(12)
                    style.paragraph_format.keep_with_next = True
                elif style_name == 'Purpose Style':
                    style.paragraph_format.left_indent = Inches(0.5)
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    style.paragraph_format.space_after = Pt(12)

        # Add document title
        if 'Title Style' in created_styles:
            title = new_doc.add_paragraph('Corporate Leadership Roles', 'Title Style')

        # Add subtitle
        if 'Subtitle Style' in created_styles:
            subtitle = new_doc.add_paragraph(
                'A Comprehensive Guide to Executive Positions', 'Subtitle Style'
            )

        # Add introduction
        if 'Purpose Style' in created_styles:
            intro = new_doc.add_paragraph(style='Purpose Style')
            intro_text = ('This document provides a detailed overview of key leadership '
                          'positions in modern organizations. Each role is described with '
                          'its primary responsibilities and strategic importance.')
            intro.add_run(intro_text).italic = True
            new_doc.add_paragraph()  # Add spacing

        # Process content
        current_role = None
        current_text: List[str] = []
        roles_summary: List[Tuple[str, str]] = []

        # Check if document has any paragraphs
        if not original_doc.paragraphs:
            console.print("[yellow]Warning: Input document is empty[/yellow]")
            return False

        for para in original_doc.paragraphs:
            is_valid, text = process_paragraph_text(para.text)
            if not is_valid:
                continue

            if is_role_title(text):
                # Add previous role if exists
                if current_role and current_text:
                    if 'Role Style' in created_styles:
                        role_para = new_doc.add_paragraph(current_role, 'Role Style')
                    if 'Purpose Style' in created_styles:
                        purpose_para = new_doc.add_paragraph(style='Purpose Style')
                        purpose_para.add_run('Purpose: ').bold = True
                        purpose_text = ' '.join(current_text)
                        purpose_para.add_run(purpose_text)
                        roles_summary.append((current_role, purpose_text))

                current_role = text
                current_text = []
            else:
                # Clean up purpose text
                if text.lower().startswith('purpose:'):
                    text = text[8:].strip()
                current_text.append(text)

        # Add last role
        if current_role and current_text:
            if 'Role Style' in created_styles:
                role_para = new_doc.add_paragraph(current_role, 'Role Style')
            if 'Purpose Style' in created_styles:
                purpose_para = new_doc.add_paragraph(style='Purpose Style')
                purpose_para.add_run('Purpose: ').bold = True
                purpose_text = ' '.join(current_text)
                purpose_para.add_run(purpose_text)
                roles_summary.append((current_role, purpose_text))

        # Add summary table if we have roles
        if roles_summary:
            new_doc.add_paragraph()  # Spacing
            if 'Heading Style' in created_styles:
                new_doc.add_paragraph('Executive Roles Summary', 'Heading Style')

            table = new_doc.add_table(rows=len(roles_summary) + 1, cols=2)
            table.autofit = False
            table.allow_autofit = False

            # Set column widths
            try:
                table.columns[0].width = Inches(2.5)
                table.columns[1].width = Inches(4.5)
            except Exception as e:
                console.print(f"[red]Warning: Could not set table column widths:[/red] {str(e)}")

            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Role'
            header_cells[1].text = 'Key Responsibilities'

            # Add data
            for i, (role, purpose) in enumerate(roles_summary, 1):
                try:
                    row_cells = table.rows[i].cells
                    row_cells[0].text = role
                    row_cells[1].text = purpose
                except Exception as e:
                    console.print(f"[red]Warning: Error adding row {i} to table:[/red] {str(e)}")
                    continue

            # Apply table styling
            create_table_style(table)

        # Save the new document
        try:
            save_task = progress.add_task("[cyan]Saving document...", total=None)
            new_doc.save(output_path)
            progress.update(save_task, completed=True)

            # Create summary table for console
            summary_table = Table(title="Document Processing Summary", show_header=True, header_style="bold magenta")
            summary_table.add_column("Item", style="cyan")
            summary_table.add_column("Status", style="green")

            summary_table.add_row("Input File", input_path)
            summary_table.add_row("Output File", output_path)
            summary_table.add_row("Roles Processed", str(len(roles_summary)))
            summary_table.add_row("Styles Applied", "✓ Title\n✓ Subtitle\n✓ Headings\n✓ Tables")

            console.print("\n")
            console.print(Panel.fit(
                "[bold green]Document processed successfully![/bold green]",
                title="Success",
                border_style="green"
            ))
            console.print(summary_table)

            return True
        except Exception as e:
            console.print(f"[red]Error saving document:[/red] {str(e)}")
            return False


if __name__ == '__main__':
    import argparse

    # Set up argument parser with rich formatting
    parser = argparse.ArgumentParser(
        description='[bold cyan]Process and style a Word document containing role descriptions.[/bold cyan]',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
[green]Examples:[/green]
  python style_roles.py input.docx
  python style_roles.py input.docx -o styled_output.docx
        '''
    )

    parser.add_argument(
        'input_file',
        help='Input Word document (.docx) to process'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output file name (default: input_file_output.docx)',
        default=None
    )

    # Parse arguments
    args = parser.parse_args()

    # Generate default output name if not provided
    if args.output is None:
        base_name = os.path.splitext(args.input_file)[0]
        args.output = f"{base_name}_output.docx"

    try:
        # Show welcome message
        console.print(Panel.fit(
            "[bold blue]Word Document Style Processor[/bold blue]\n"
            "[cyan]A tool to enhance and style your Word documents[/cyan]",
            border_style="blue"
        ))

        success = clean_and_style_document(args.input_file, args.output)
        if not success:
            console.print(Panel.fit(
                "[bold red]Failed to process document[/bold red]\n"
                "Please check the error messages above.",
                title="Error",
                border_style="red"
            ))

    except KeyboardInterrupt:
        console.print("\n[yellow]Process interrupted by user[/yellow]")
    except Exception as e:
        console.print(f"\n[red]Unexpected error:[/red] {str(e)}")
        console.print("\n[yellow]Usage help:[/yellow]")
        parser.print_help()
